import base64
import io
from unittest.mock import MagicMock, patch

import pytest
from fastapi.testclient import TestClient
from pypdf import PdfWriter
from pypdf.generic import DecodedStreamObject, DictionaryObject, NameObject

from backend.document_content import (
    INVALID_DOCUMENT_CODE,
    DocumentContentError,
    extract_document_text,
)
from backend.main import app


_client = TestClient(app, raise_server_exceptions=False)


def _blank_pdf() -> bytes:
    output = io.BytesIO()
    writer = PdfWriter()
    writer.add_blank_page(width=72, height=72)
    writer.write(output)
    return output.getvalue()


def _pdf_with_sensitive_broken_cmap(sentinel: str) -> bytes:
    """파싱은 성공하지만 pypdf가 invalid CMap token을 warning에 넣는 PDF."""
    output = io.BytesIO()
    writer = PdfWriter()
    page = writer.add_blank_page(width=72, height=72)

    cmap = DecodedStreamObject()
    cmap.set_data(
        (
            "/CIDInit /ProcSet findresource begin\n"
            "12 dict begin\n"
            "begincmap\n"
            "/CMapType 2 def\n"
            "1 begincodespacerange\n<00> <FF>\nendcodespacerange\n"
            "2 beginbfchar\n"
            "<41> <0041>\n"
            f"<42> <{sentinel}>\n"
            "endbfchar\nendcmap\nend\nend\n"
        ).encode()
    )
    font = DictionaryObject({
        NameObject("/Type"): NameObject("/Font"),
        NameObject("/Subtype"): NameObject("/Type1"),
        NameObject("/BaseFont"): NameObject("/Helvetica"),
        NameObject("/Encoding"): NameObject("/WinAnsiEncoding"),
        NameObject("/ToUnicode"): writer._add_object(cmap),
    })
    resources = DictionaryObject({
        NameObject("/Font"): DictionaryObject({
            NameObject("/F1"): writer._add_object(font),
        }),
    })
    content = DecodedStreamObject()
    content.set_data(b"BT /F1 12 Tf 10 10 Td <41> Tj ET")
    page[NameObject("/Resources")] = resources
    page[NameObject("/Contents")] = writer._add_object(content)
    writer.write(output)
    return output.getvalue()


DOCUMENT_CASES = [
    pytest.param("utf8.txt", "한글 문서".encode(), True, "한글 문서", id="utf8-korean"),
    pytest.param("bom.txt", b"\xef\xbb\xbfhello", True, "hello", id="utf8-bom"),
    pytest.param("crlf.md", b"line1\r\nline2", True, "line1\r\nline2", id="crlf"),
    pytest.param("empty.txt", b"", True, "", id="empty"),
    pytest.param("cp949.txt", "한글 문서".encode("cp949"), True, "한글 문서", id="cp949"),
    pytest.param("invalid.txt", b"\x81", False, None, id="undecodable"),
    pytest.param("nul.txt", b"hello\x00world", False, None, id="nul"),
    pytest.param("control-ok.txt", b"\x01\x02" + b"a" * 98, True, None, id="control-2pct"),
    pytest.param("control-bad.txt", b"\x01\x02\x03" + b"a" * 97, False, None, id="control-over-2pct"),
    pytest.param("normal.pdf", _blank_pdf(), True, "", id="normal-pdf"),
    pytest.param("fake.pdf", b"not-a-pdf", False, None, id="pdf-mismatch"),
    pytest.param("broken.pdf", b"%PDF-not-parseable", False, None, id="pdf-header-only"),
]


# 매직 검사는 통과하지만 파서가 실패하는 입력. 업로드·질의 HTTP 경로에서는 입력 경계
# 위반(415)이 아니라 변환 실패로 분류된다 — 업로드 400, 질의는 관대한 placeholder.
# extract_document_text() 직접 호출 계약은 그대로 DocumentContentError다.
CONVERSION_FAILURE_FILENAMES = {"broken.pdf"}


@pytest.mark.parametrize(("filename", "data", "valid", "expected"), DOCUMENT_CASES)
def test_shared_document_validator_fixture_matrix(filename, data, valid, expected):
    if not valid:
        with pytest.raises(DocumentContentError) as exc_info:
            extract_document_text(filename, data)
        assert exc_info.value.code == INVALID_DOCUMENT_CODE
        return

    actual = extract_document_text(filename, data)
    if expected is not None:
        assert actual == expected


def _project_conn():
    conn = MagicMock()
    cursor = conn.cursor.return_value.__enter__.return_value
    cursor.fetchone.return_value = {"id": 1}
    cursor.fetchall.return_value = []
    cursor.lastrowid = 7
    return conn


def _reservation():
    return {
        "reservation_id": "reservation-1",
        "temp_path": "/tmp/document-fixture.tmp",
        "target_path": "/tmp/document-fixture",
    }


def _finalized():
    return {
        "doc_id": 7,
        "old_doc_ids": [],
        "file_path": "/tmp/document-fixture",
        "processing_token": "token-1",
    }


@pytest.mark.parametrize(("filename", "data", "valid", "expected"), DOCUMENT_CASES)
def test_upload_applies_shared_fixture_before_db_and_storage(filename, data, valid, expected):
    with patch("backend.api.documents.require_project_access"), patch(
        "backend.api.documents.require_upload_user", return_value=1
    ), patch(
        "backend.api.documents.reserve_document", return_value=_reservation()
    ) as reserve, patch(
        "backend.api.documents.write_reserved_file"
    ) as write_file, patch(
        "backend.api.documents.finalize_document", return_value=_finalized()
    ), patch("backend.api.documents._process_upload") as process:
        response = _client.post(
            "/api/v1/projects/1/documents",
            files={"file": (filename, data, "application/octet-stream")},
        )

    if filename in CONVERSION_FAILURE_FILENAMES:
        # 매직 통과 후 변환 실패 → 400 + 세부 code(문자열 detail). DB/storage는 건드리지 않는다.
        assert response.status_code == 400
        body = response.json()
        assert isinstance(body["detail"], str)
        assert body["code"] == "corrupt_file"
        reserve.assert_not_called()
        write_file.assert_not_called()
        process.assert_not_called()
    elif not valid:
        assert response.status_code == 415
        assert response.json()["detail"]["code"] == INVALID_DOCUMENT_CODE
        reserve.assert_not_called()
        write_file.assert_not_called()
        process.assert_not_called()
    elif expected == "":
        # upload의 기존 empty semantics는 400이며 DB/storage보다 먼저 끝난다.
        assert response.status_code == 400
        reserve.assert_not_called()
        write_file.assert_not_called()
        process.assert_not_called()
    else:
        assert response.status_code == 201
        reserve.assert_called_once()
        write_file.assert_called_once()
        process.assert_called_once()


@pytest.mark.parametrize(("filename", "data", "valid", "expected"), DOCUMENT_CASES)
def test_query_applies_same_fixture_before_db_and_llm(filename, data, valid, expected):
    conn = _project_conn()
    payload = {
        "question": "fixture test",
        "attachments": [{
            "filename": filename,
            "content_base64": base64.b64encode(data).decode(),
        }],
    }
    with patch("backend.api.query.require_project_access"), patch(
        "backend.api.query.get_connection", return_value=conn
    ) as get_connection, patch(
        "backend.api.query.run_agentic_qa", return_value={"answer": "ok", "debug": {}}
    ) as run_agentic_qa:
        response = _client.post("/api/v1/projects/1/query", json=payload)

    if filename in CONVERSION_FAILURE_FILENAMES:
        # 검증을 통과한 파일의 변환 실패는 질의를 막지 않는다(관대한 첨부 정책).
        assert response.status_code == 200
        assert get_connection.called
        run_agentic_qa.assert_called_once()
    elif not valid:
        assert response.status_code == 415
        assert response.json()["detail"]["code"] == INVALID_DOCUMENT_CODE
        get_connection.assert_not_called()
        run_agentic_qa.assert_not_called()
    else:
        # query는 empty attachment를 명시적 placeholder로 유지하는 기존 계약이다.
        assert response.status_code == 200
        assert get_connection.called
        run_agentic_qa.assert_called_once()


def test_sensitive_pdf_parser_warning_is_absent_from_shared_and_endpoint_logs(caplog):
    sentinel = "SENSITIVE-SENTINEL"
    data = _pdf_with_sensitive_broken_cmap(sentinel)
    assert sentinel.encode() in data

    caplog.clear()
    assert extract_document_text("sensitive.pdf", data) == "A"

    with patch("backend.api.documents.require_project_access"), patch(
        "backend.api.documents.require_upload_user", return_value=1
    ), patch("backend.api.documents.reserve_document", return_value=_reservation()), patch(
        "backend.api.documents.write_reserved_file"
    ), patch("backend.api.documents.finalize_document", return_value=_finalized()), patch(
        "backend.api.documents._process_upload"
    ):
        upload_response = _client.post(
            "/api/v1/projects/1/documents",
            files={"file": ("sensitive.pdf", data, "application/pdf")},
        )

    query_conn = _project_conn()
    with patch("backend.api.query.require_project_access"), patch(
        "backend.api.query.get_connection", return_value=query_conn
    ), patch(
        "backend.api.query.run_agentic_qa", return_value={"answer": "ok", "debug": {}}
    ):
        query_response = _client.post(
            "/api/v1/projects/1/query",
            json={
                "question": "fixture test",
                "attachments": [{
                    "filename": "sensitive.pdf",
                    "content_base64": base64.b64encode(data).decode(),
                }],
            },
        )

    assert upload_response.status_code == 201
    assert query_response.status_code == 200
    assert sentinel not in upload_response.text
    assert sentinel not in query_response.text
    assert sentinel not in caplog.text


# ─── 파싱 중복 제거 회귀 (C안) ────────────────────────────────────────────────
#
# document_content 의 DOCX·PDF 파서는 pipeline/converters 에 위임한다. 예전에는
# 여기에 별도 구현이 있었는데 중첩 표를 최상위만 읽어 셀 안의 표를 통째로 잃었다.
# 두 구현이 다시 갈라지면 같은 파일이 호출 경로에 따라 다르게 처리되므로,
# 아래 테스트들이 위임이 풀리는 순간 실패하도록 고정한다.

def _docx_with_nested_table() -> bytes:
    import docx

    document = docx.Document()
    document.add_paragraph("예산 검토")
    outer = document.add_table(rows=1, cols=2)
    outer.cell(0, 0).text = "예산"
    inner_host = outer.cell(0, 1)
    inner_host.text = "내역:"
    inner = inner_host.add_table(rows=1, cols=2)
    inner.cell(0, 0).text = "인건비"
    inner.cell(0, 1).text = "3000만원"
    buffer = io.BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def _multipage_pdf(pages: list[list[str]]) -> bytes:
    """외부 렌더러 없이 텍스트 레이어가 있는 PDF를 조립한다(테스트 공용 헬퍼 재사용)."""
    from tests.test_converters import _make_pdf

    return _make_pdf(pages)


def test_registry_matches_converter_registry():
    """두 레지스트리는 별도 객체다. 드리프트하면 업로드·질의 지원 포맷이 어긋난다."""
    from backend.document_content import DOCUMENT_PARSERS
    from backend.pipeline.converters import supported_suffixes

    assert set(DOCUMENT_PARSERS) == set(supported_suffixes())


def test_docx_extraction_delegates_to_converter():
    """DOCX 직접 추출 결과가 변환기 결과와 같아야 한다 — 위임이 풀리면 실패한다."""
    from backend.pipeline.converters import convert

    data = _docx_with_nested_table()
    assert extract_document_text("x.docx", data) == convert("x.docx", data).text
    # 중첩 표 내용이 살아 있는지 직접 확인한다. 위임 전 구현은 이걸 잃었다.
    assert "인건비" in extract_document_text("x.docx", data)


def test_pdf_extraction_delegates_to_converter():
    """PDF도 동일하다. DOCX 테스트만으로는 _read_pdf 회귀를 잡지 못한다."""
    from backend.pipeline.converters import convert

    data = _multipage_pdf([["First page body text."], ["Second page body text."]])
    assert extract_document_text("x.pdf", data) == convert("x.pdf", data).text


def test_empty_docx_returns_empty_string():
    """빈 DOCX는 EMPTY_DOCUMENT 분기를 타고 "" 로 복원된다."""
    import docx

    buffer = io.BytesIO()
    docx.Document().save(buffer)
    assert extract_document_text("empty.docx", buffer.getvalue()) == ""


def test_pdf_without_text_layer_returns_empty_string():
    """텍스트 레이어가 없는 PDF는 NO_TEXT_LAYER 분기를 탄다."""
    assert extract_document_text("blank.pdf", _blank_pdf()) == ""


def test_pdf_with_only_noise_returns_empty_string():
    """페이지 번호·구분선만 있는 PDF는 정제 후 블록이 비어 EMPTY_DOCUMENT를 탄다.

    빈 페이지(NO_TEXT_LAYER)와 다른 분기이므로 별도로 고정한다.
    """
    data = _multipage_pdf([["- 1 -"], ["- 2 -"], ["- 3 -"]])
    assert extract_document_text("noise.pdf", data) == ""


def test_oversized_docx_keeps_actionable_message():
    """한도 초과 메시지의 조치 안내가 어댑터를 거쳐도 남아야 한다.

    예외 타입만 단언하면 "올바른 DOCX가 아닙니다" 같은 일반 문구로 덮여도 통과한다.
    사용자가 무엇을 해야 하는지가 사라지는 회귀를 잡는다.
    """
    import docx

    from backend.pipeline.converters import docx_converter

    buffer = io.BytesIO()
    document = docx.Document()
    document.add_paragraph("본문")
    document.save(buffer)

    original = docx_converter._MAX_COMPRESSION_RATIO
    docx_converter._MAX_COMPRESSION_RATIO = 0  # 큰 파일을 만들지 않고 한도만 낮춘다
    try:
        with pytest.raises(DocumentContentError) as exc_info:
            extract_document_text("big.docx", buffer.getvalue())
    finally:
        docx_converter._MAX_COMPRESSION_RATIO = original

    assert exc_info.value.code == INVALID_DOCUMENT_CODE
    assert "나누거나" in exc_info.value.message
    assert "줄여서" in exc_info.value.message


def test_text_formats_keep_original_contract():
    """MD/TXT는 위임 대상이 아니다. CRLF·BOM·CP949·인코딩 계약이 그대로여야 한다."""
    assert extract_document_text("crlf.md", b"line1\r\nline2") == "line1\r\nline2"
    assert extract_document_text("bom.txt", b"\xef\xbb\xbfhello") == "hello"
    assert extract_document_text("cp949.txt", "한글".encode("cp949")) == "한글"
    with pytest.raises(DocumentContentError) as exc_info:
        extract_document_text("bad.txt", b"\x81")
    assert exc_info.value.code == INVALID_DOCUMENT_CODE
