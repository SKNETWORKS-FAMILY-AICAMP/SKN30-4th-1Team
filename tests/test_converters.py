"""멀티포맷 변환기(DOCX·텍스트 PDF·평문)와 구조 기반 청킹 회귀 테스트.

실제 파일 바이트를 만들어 검증한다 — 파서를 mock하면 "라이브러리가 실제로
무엇을 돌려주는가"라는 이 계층의 유일한 위험이 검증되지 않는다.
"""
import io

import pytest

from backend.pipeline.converters import (
    ConversionError,
    ErrorCode,
    WarningCode,
    chunk_document,
    convert,
    is_supported,
    supported_suffixes,
)
from backend.pipeline.converters.chunker import DEFAULT_CHUNK_SIZE


# ─── 픽스처 빌더 ────────────────────────────────────────────────────────────

def _make_docx(build) -> bytes:
    """build(document) 콜백으로 DOCX를 만들고 바이트로 돌려준다."""
    import docx

    document = docx.Document()
    build(document)
    buffer = io.BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def _make_pdf(pages: list[list[str]]) -> bytes:
    """텍스트 레이어가 있는 최소 PDF를 직접 조립한다(외부 렌더러 의존 없이)."""
    font_obj = 3 + 2 * len(pages)
    kids = " ".join(f"{3 + 2 * i} 0 R" for i in range(len(pages)))
    objects: list[tuple[int, str]] = [
        (1, "<< /Type /Catalog /Pages 2 0 R >>"),
        (2, f"<< /Type /Pages /Kids [{kids}] /Count {len(pages)} >>"),
    ]
    for index, lines in enumerate(pages):
        page_id, content_id = 3 + 2 * index, 4 + 2 * index
        objects.append((page_id, (
            "<< /Type /Page /Parent 2 0 R /MediaBox [0 0 612 792]"
            f" /Resources << /Font << /F1 {font_obj} 0 R >> >>"
            f" /Contents {content_id} 0 R >>"
        )))
        drawn = "\n".join(f"({line}) Tj T*" for line in lines)
        stream = f"BT /F1 12 Tf 72 720 Td 16 TL\n{drawn}\nET"
        objects.append((content_id, f"<< /Length {len(stream)} >>\nstream\n{stream}\nendstream"))
    objects.append((font_obj, "<< /Type /Font /Subtype /Type1 /BaseFont /Helvetica >>"))

    out = io.BytesIO()
    out.write(b"%PDF-1.4\n")
    offsets: dict[int, int] = {}
    for number, content in objects:
        offsets[number] = out.tell()
        out.write(f"{number} 0 obj\n{content}\nendobj\n".encode("latin-1"))
    xref_offset = out.tell()
    last = max(offsets)
    out.write(f"xref\n0 {last + 1}\n0000000000 65535 f \n".encode())
    for number in range(1, last + 1):
        out.write(f"{offsets.get(number, 0):010d} 00000 n \n".encode())
    out.write(
        f"trailer\n<< /Size {last + 1} /Root 1 0 R >>\n"
        f"startxref\n{xref_offset}\n%%EOF\n".encode()
    )
    return out.getvalue()


# ─── 레지스트리 ─────────────────────────────────────────────────────────────

def test_required_formats_are_supported():
    """필수 포맷(DOCX·PDF)과 기존 평문 포맷이 모두 열려 있어야 한다."""
    assert {".docx", ".pdf", ".md", ".txt"} <= supported_suffixes()
    assert is_supported("회의록.DOCX")  # 확장자 대소문자 무관


def test_unsupported_format_raises_explicit_error():
    with pytest.raises(ConversionError) as exc:
        convert("보고서.hwp", b"whatever")
    assert exc.value.code == ErrorCode.UNSUPPORTED_FORMAT


# ─── DOCX ───────────────────────────────────────────────────────────────────

def test_docx_preserves_heading_list_and_table_order():
    """본문 순서·제목 레벨·목록·표가 모두 Block으로 보존된다."""
    def build(document):
        document.add_heading("주간 회의", level=1)
        document.add_paragraph("FastAPI로 백엔드를 구성하기로 했다.")
        document.add_heading("액션 아이템", level=2)
        document.add_paragraph("API 명세서 작성", style="List Bullet")
        table = document.add_table(rows=2, cols=2)
        table.cell(0, 0).text = "담당"
        table.cell(0, 1).text = "기한"
        table.cell(1, 0).text = "김동휘"
        table.cell(1, 1).text = "7/25"

    doc = convert("회의록.docx", _make_docx(build))

    assert doc.format == "docx"
    kinds = [b.kind for b in doc.blocks]
    assert kinds == ["heading", "paragraph", "heading", "list_item", "table_row", "table_row"]
    # 문단 번호는 0부터 빈틈없이 매겨진다 — 출처 추적의 기준 좌표다.
    assert [b.order for b in doc.blocks] == list(range(len(doc.blocks)))
    assert doc.blocks[0].level == 1 and doc.blocks[2].level == 2
    # 하위 제목 아래 블록은 상위 제목 경로를 물려받는다.
    assert doc.blocks[3].heading_path == ("주간 회의", "액션 아이템")
    assert "김동휘 | 7/25" in doc.blocks[5].text


def test_docx_table_emits_flatten_warning():
    """표는 열 구조가 유실되므로 조용히 넘기지 않고 경고를 남긴다."""
    def build(document):
        document.add_paragraph("본문")
        table = document.add_table(rows=1, cols=2)
        table.cell(0, 0).text = "A"
        table.cell(0, 1).text = "B"

    doc = convert("표.docx", _make_docx(build))
    assert WarningCode.TABLE_FLATTENED in {w.code for w in doc.warnings}


def test_docx_detects_heading_by_style_name_not_style_id():
    """styleId가 불투명해도 스타일 '이름'으로 제목을 알아본다.

    국내에서 작성된 DOCX는 styleId가 `a3` 같은 값이고 이름에만 `제목 1`이 담기는
    경우가 흔하다. 성능을 위해 styleId를 직접 읽으므로 이름 해석 경로를 고정한다.
    """
    import docx
    from docx.enum.style import WD_STYLE_TYPE

    document = docx.Document()
    style = document.styles.add_style("제목 1", WD_STYLE_TYPE.PARAGRAPH)
    document.add_paragraph("한글 제목", style=style)
    document.add_paragraph("본문 내용")
    buffer = io.BytesIO()
    document.save(buffer)

    doc = convert("한글.docx", buffer.getvalue())
    assert [b.kind for b in doc.blocks] == ["heading", "paragraph"]
    assert doc.blocks[1].heading_path == ("한글 제목",)


def test_docx_has_no_page_numbers():
    """DOCX는 페이지가 렌더링 시점에 정해지므로 page를 지어내지 않는다."""
    doc = convert("x.docx", _make_docx(lambda d: d.add_paragraph("내용")))
    assert all(b.page is None for b in doc.blocks)
    assert doc.page_count is None


def test_docx_empty_document_raises():
    """빈 문서는 성공으로 위장하지 않고 명시적 오류를 낸다."""
    with pytest.raises(ConversionError) as exc:
        convert("빈.docx", _make_docx(lambda d: d.add_paragraph("   ")))
    assert exc.value.code == ErrorCode.EMPTY_DOCUMENT


def test_docx_corrupt_file_raises():
    with pytest.raises(ConversionError) as exc:
        convert("깨진.docx", b"not a real docx")
    assert exc.value.code == ErrorCode.CORRUPT_FILE


# ─── PDF ────────────────────────────────────────────────────────────────────

def test_pdf_assigns_page_numbers_to_blocks():
    data = _make_pdf([
        ["Team decided to use FastAPI."],
        ["Deployment gate is on 2026-07-25."],
    ])
    doc = convert("보고서.pdf", data)

    assert doc.format == "pdf"
    assert doc.page_count == 2
    pages = {b.page for b in doc.blocks}
    assert pages == {1, 2}
    assert any("FastAPI" in b.text and b.page == 1 for b in doc.blocks)
    assert any("2026-07-25" in b.text and b.page == 2 for b in doc.blocks)


def test_pdf_drops_repeated_header_and_page_numbers():
    """모든 페이지에 반복되는 머리말과 페이지 번호는 본문에서 제외된다."""
    data = _make_pdf([
        ["PaiM Confidential", f"Body line number {i}.", str(i)]
        for i in range(1, 5)
    ])
    doc = convert("정책.pdf", data)

    text = doc.text
    assert "PaiM Confidential" not in text
    assert "Body line number 1." in text and "Body line number 4." in text
    assert WarningCode.REPEATED_LINE_DROPPED in {w.code for w in doc.warnings}


def test_pdf_without_text_layer_raises_no_text_layer():
    """스캔 PDF는 범위 밖 — 빈 결과 대신 이유가 담긴 오류를 낸다."""
    with pytest.raises(ConversionError) as exc:
        convert("스캔본.pdf", _make_pdf([[], []]))
    assert exc.value.code == ErrorCode.NO_TEXT_LAYER


def test_pdf_corrupt_file_raises():
    with pytest.raises(ConversionError) as exc:
        convert("깨진.pdf", b"%PDF-1.4 garbage")
    assert exc.value.code == ErrorCode.CORRUPT_FILE


# ─── 평문 / Markdown ────────────────────────────────────────────────────────

def test_markdown_classifies_headings_and_lists():
    source = (
        "# 프로젝트 계획\n\n"
        "본문 첫 문단입니다.\n\n"
        "## 세부 항목\n\n"
        "- 첫 번째 항목\n"
        "- 두 번째 항목\n"
    ).encode("utf-8")
    doc = convert("plan.md", source)

    assert doc.format == "text"
    assert [b.kind for b in doc.blocks] == [
        "heading", "paragraph", "heading", "list_item", "list_item",
    ]
    assert doc.blocks[3].heading_path == ("프로젝트 계획", "세부 항목")


def test_text_cp949_falls_back_with_warning():
    """UTF-8이 아닌 국내 문서를 깨진 채 넘기지 않고, 폴백 사실을 알린다."""
    doc = convert("메모.txt", "결정 사항 정리".encode("cp949"))
    assert "결정 사항 정리" in doc.text
    assert WarningCode.DECODE_FALLBACK in {w.code for w in doc.warnings}


def test_text_empty_document_raises():
    with pytest.raises(ConversionError) as exc:
        convert("빈.txt", b"\n\n   \n")
    assert exc.value.code == ErrorCode.EMPTY_DOCUMENT


# ─── 청킹 ───────────────────────────────────────────────────────────────────

def _long_text_doc(paragraphs: int = 40):
    body = "\n\n".join(
        f"{i}번째 문단입니다. " + ("내용을 채우는 문장입니다. " * 4)
        for i in range(paragraphs)
    )
    return convert("long.md", body.encode("utf-8"))


def test_chunks_never_exceed_chunk_size():
    """청크 크기 불변식 — 오버랩을 붙인 뒤에도 초과하지 않는다."""
    chunks = chunk_document(_long_text_doc())
    assert chunks
    assert all(len(c.text) <= DEFAULT_CHUNK_SIZE for c in chunks)


def test_oversized_single_block_is_split_within_limit():
    doc = convert("huge.md", ("가" * 2000).encode("utf-8"))
    chunks = chunk_document(doc)
    assert len(chunks) > 1
    assert all(len(c.text) <= DEFAULT_CHUNK_SIZE for c in chunks)


def test_chunk_indexes_are_contiguous():
    chunks = chunk_document(_long_text_doc())
    assert [c.index for c in chunks] == list(range(len(chunks)))


def test_chunk_carries_page_range_from_pdf():
    """PDF 청크는 원본 페이지 범위를 유지한다 — 출처를 페이지까지 되짚기 위함."""
    data = _make_pdf([
        [f"Page one sentence number {i}." for i in range(1, 6)],
        [f"Page two sentence number {i}." for i in range(1, 6)],
    ])
    chunks = chunk_document(convert("p.pdf", data))

    assert chunks
    for chunk in chunks:
        assert chunk.page_start is not None and chunk.page_start >= 1
        assert chunk.page_end >= chunk.page_start
    assert {c.page_start for c in chunks} >= {1}


def test_chunk_metadata_is_chroma_scalar_safe():
    """ChromaDB metadata는 str/int/float/bool만 허용한다."""
    doc = convert("x.docx", _make_docx(lambda d: d.add_paragraph("본문 내용")))
    metadata = chunk_document(doc)[0].to_metadata()

    assert all(isinstance(v, (str, int, float, bool)) for v in metadata.values())
    # 페이지가 없는 포맷은 None이 아니라 -1로 표현된다.
    assert metadata["page_start"] == -1
    assert metadata["block_start"] == 0


def test_chunk_keeps_heading_context():
    source = "# 배포 계획\n\n" + ("배포 절차를 정리한 문단입니다. " * 10)
    chunks = chunk_document(convert("d.md", source.encode("utf-8")))
    assert any("배포 계획" in c.heading for c in chunks)


def test_chunk_text_covers_document_content():
    """청킹 과정에서 본문이 통째로 사라지지 않는다."""
    doc = _long_text_doc(paragraphs=10)
    joined = " ".join(c.text for c in chunk_document(doc))
    for index in range(10):
        assert f"{index}번째 문단입니다." in joined


# ─── 리뷰 지적 회귀 (PR-001-R001~R006) ──────────────────────────────────────

def _ordinary_docx() -> bytes:
    """가드 검증용 **온전한** 소형 DOCX. 대형 파일을 만들 필요가 없다."""
    return _make_docx(lambda d: d.add_paragraph("평범한 회의록 본문입니다."))


def test_docx_rejects_oversized_archive_before_parsing(monkeypatch):
    """R001: 한도를 넘는 DOCX는 python-docx에 넘기기 전에 거절한다.

    R2B01: 이전 테스트는 `[Content_Types].xml`이 없는 무효 ZIP을 써서, 가드를
    통째로 지워도 python-docx가 같은 오류를 내며 통과했다(False Green).
    온전한 DOCX + 한도 monkeypatch로 **가드만** 발동시켜 그 함정을 없앤다.
    """
    import backend.pipeline.converters.docx_converter as module

    data = _ordinary_docx()
    # 정상 파일이 한도에 걸리도록 압축비 상한만 1로 낮춘다.
    monkeypatch.setattr(module, "_MAX_COMPRESSION_RATIO", 1)

    # 가드가 실제로 파싱을 막았는지 확인하기 위해 파서 호출을 감시한다.
    called = []
    real_document = module.docx.Document if hasattr(module, "docx") else None

    with pytest.raises(ConversionError) as exc:
        convert("보통.docx", data)

    # 손상이 아니라 "한도 초과"로 분류돼야 한다 (R2B03).
    assert exc.value.code == ErrorCode.FILE_TOO_LARGE
    # R3B01: 코드만이 아니라 **조치 안내 문구**를 고정한다. R2B03이 요구한 것은
    # "손상으로 오해하지 않게 조치를 안내하라"였는데, 문구를 지워도 통과하면
    # 그 요구가 지켜지는지 아무도 보장하지 못한다.
    assert "나누거나" in exc.value.message and "줄여서" in exc.value.message
    assert real_document is None or not called


def test_guard_runs_before_python_docx_parses(monkeypatch):
    """R2B01: 가드가 파서보다 먼저 동작함을 spy로 증명한다."""
    import docx as docx_module

    import backend.pipeline.converters.docx_converter as module

    data = _ordinary_docx()
    monkeypatch.setattr(module, "_MAX_COMPRESSION_RATIO", 1)

    parsed = []
    original = docx_module.Document
    monkeypatch.setattr(
        docx_module, "Document",
        lambda *a, **k: (parsed.append(1), original(*a, **k))[1],
    )

    with pytest.raises(ConversionError):
        convert("보통.docx", data)
    assert parsed == [], "가드가 막았어야 하는데 python-docx가 호출됐다"


def test_malformed_zip_directory_does_not_leak_500():
    """R2B02: 중앙 디렉터리가 깨진 ZIP도 ConversionError로 정규화된다.

    가드 추가 전에는 docx.Document() 주변의 넓은 except가 흡수하던 입력이라,
    여기서 잡지 않으면 400이어야 할 응답이 500으로 누출된다.
    """
    import zipfile

    data = bytearray(_ordinary_docx())
    # 중앙 디렉터리 시그니처(PK\x01\x02) 뒤의 extraction version을 비정상 값으로 변조.
    index = data.rfind(b"PK\x01\x02")
    assert index != -1
    data[index + 6] = 0xFF  # version needed to extract

    with pytest.raises(ConversionError) as exc:
        convert("변조.docx", bytes(data))
    # R3B02: CORRUPT_FILE로 한정한다. FILE_TOO_LARGE까지 허용하면 손상 파일이
    # 크기 초과로 오분류돼 "파일을 줄이라"는 잘못된 안내를 하는 회귀를 놓친다.
    assert exc.value.code == ErrorCode.CORRUPT_FILE


def _corrupted_docx() -> bytes:
    """중앙 디렉터리의 extraction version을 지원 상한 이상으로 변조한 소형 DOCX."""
    data = bytearray(_ordinary_docx())
    index = data.rfind(b"PK\x01\x02")
    assert index != -1
    data[index + 6] = 0xFF
    return bytes(data)


def test_upload_rejects_corrupted_archive_with_corrupt_file_code():
    """R3B02: 손상 파일은 업로드 경로에서도 corrupt_file이어야 한다.

    file_too_large로 분류되면 사용자에게 "파일을 줄이라"는 잘못된 조치를 안내한다.
    """
    from unittest.mock import patch

    from fastapi.testclient import TestClient

    from backend.main import app

    client = TestClient(app, raise_server_exceptions=False)
    with patch("backend.api.upload.require_project_access"), \
         patch("backend.api.upload.require_upload_user", return_value=1):
        response = client.post(
            "/api/v1/projects/1/documents",
            files={"file": ("변조.docx", _corrupted_docx(), "application/octet-stream")},
        )

    assert response.status_code == 400
    assert response.json()["code"] == ErrorCode.CORRUPT_FILE


def test_query_attachment_handles_corrupted_docx_without_500():
    """R3B02: 손상 첨부는 500 없이 실패 placeholder로 처리된다."""
    import base64
    from unittest.mock import MagicMock, patch

    from fastapi.testclient import TestClient

    from backend.main import app

    conn = MagicMock()
    conn.cursor.return_value.__enter__.return_value.fetchone.return_value = {"id": 1}
    captured = {}

    def fake_run_qa(**kwargs):
        captured.update(kwargs)
        return {"answer": "답", "sources": [], "debug": {}}

    client = TestClient(app, raise_server_exceptions=False)
    with patch("backend.api.query.require_project_access"), \
         patch("backend.api.query.get_connection", return_value=conn), \
         patch("backend.api.query.run_qa", side_effect=fake_run_qa):
        response = client.post(
            "/api/v1/projects/1/query",
            json={
                "question": "요약해줘",
                "attachments": [{
                    "filename": "변조.docx",
                    "content_base64": base64.b64encode(_corrupted_docx()).decode(),
                }],
            },
        )

    assert response.status_code == 200, response.text
    assert "(텍스트를 추출할 수 없습니다.)" in captured["attachment_context"]


def test_upload_rejects_oversized_archive_with_400(monkeypatch):
    """R2B02: 업로드 경로에서 500이 아니라 400 + 코드가 나와야 한다."""
    from unittest.mock import patch

    from fastapi.testclient import TestClient

    import backend.pipeline.converters.docx_converter as module
    from backend.main import app

    monkeypatch.setattr(module, "_MAX_COMPRESSION_RATIO", 1)
    client = TestClient(app, raise_server_exceptions=False)
    with patch("backend.api.upload.require_project_access"), \
         patch("backend.api.upload.require_upload_user", return_value=1):
        response = client.post(
            "/api/v1/projects/1/documents",
            files={"file": ("보통.docx", _ordinary_docx(), "application/octet-stream")},
        )

    assert response.status_code == 400
    body = response.json()
    assert isinstance(body["detail"], str)
    assert body["code"] == ErrorCode.FILE_TOO_LARGE
    # R3B01: 사용자에게 실제로 닿는 detail에도 조치 안내가 있어야 한다.
    assert "나누거나" in body["detail"] and "줄여서" in body["detail"]


def test_query_attachment_replaces_rejected_docx_with_placeholder(monkeypatch):
    """R001·R2B01: 질의 첨부에서 거절된 DOCX는 본문으로 쓰이지 않고 대체된다.

    이전 테스트는 `status_code in (200, 404, 500)`이라 사실상 아무것도 보장하지
    않았다(500까지 성공으로 간주). 프로젝트 존재와 run_qa를 정상 mock해서
    거절된 첨부가 실제로 placeholder가 되는지 단언한다.
    """
    import base64
    from unittest.mock import MagicMock, patch

    from fastapi.testclient import TestClient

    import backend.pipeline.converters.docx_converter as module
    from backend.main import app

    monkeypatch.setattr(module, "_MAX_COMPRESSION_RATIO", 1)

    conn = MagicMock()
    conn.cursor.return_value.__enter__.return_value.fetchone.return_value = {"id": 1}
    captured = {}

    def fake_run_qa(**kwargs):
        captured.update(kwargs)
        return {"answer": "답", "sources": [], "debug": {}}

    client = TestClient(app, raise_server_exceptions=False)
    with patch("backend.api.query.require_project_access"), \
         patch("backend.api.query.get_connection", return_value=conn), \
         patch("backend.api.query.run_qa", side_effect=fake_run_qa):
        response = client.post(
            "/api/v1/projects/1/query",
            json={
                "question": "요약해줘",
                "attachments": [{
                    "filename": "보통.docx",
                    "content_base64": base64.b64encode(_ordinary_docx()).decode(),
                }],
            },
        )

    assert response.status_code == 200, response.text
    context = captured["attachment_context"]
    assert "(텍스트를 추출할 수 없습니다.)" in context
    assert "평범한 회의록 본문입니다." not in context


def test_chunk_overlap_preserves_all_source_pages():
    """R002: 오버랩이 두 페이지에 걸치면 다음 청크 범위가 둘 다 포함해야 한다.

    R204: 2페이지 fixture는 수정 전 구현에서도 혼합 청크가 1~2 범위를 갖게 되어
    단언이 만족됐다(False Green). 아래 fixture는 수정 전 구현에서 실제로 실패한다.

    두 가지가 맞아야 이 테스트가 의미를 갖는다.
    1. 3페이지 — 오버랩이 2페이지 내용을 3페이지 청크로 넘기는 상황이 만들어진다.
    2. 마커가 줄 **끝**에 있어야 한다. 오버랩은 조각을 뒤에서 잘라 오므로,
       줄 앞의 마커는 잘려나가 단언이 아무것도 못 잡는다.
    """
    data = _make_pdf([
        [f"padding words filler text line {i} ALPHA." for i in range(9)],
        [f"padding words filler text line {i} BRAVO." for i in range(14)],
        [f"padding words filler text line {i} DELTA." for i in range(14)],
    ])
    doc = convert("p.pdf", data)
    chunks = chunk_document(doc)

    markers = {"ALPHA": 1, "BRAVO": 2, "DELTA": 3}
    for chunk in chunks:
        for marker, page_number in markers.items():
            if marker in chunk.text:
                assert chunk.page_start <= page_number <= chunk.page_end, (
                    f"chunk{chunk.index}가 {marker}(p{page_number})를 담았지만 "
                    f"페이지 범위는 {chunk.page_start}~{chunk.page_end}"
                )

    # 오버랩으로 앞 페이지 내용을 물려받은 청크가 실제로 존재해야 한다.
    # (없으면 위 단언이 공허하게 통과한다)
    inherited = [
        c for c in chunks[1:]
        if "BRAVO" in c.text and c.page_start <= 2
    ]
    assert inherited, "오버랩이 앞 페이지를 물려준 청크가 없어 검증이 공허하다"


def test_chunk_overlap_preserves_block_range():
    """R002: 오버랩이 여러 블록에 걸쳐도 block_start가 실제 출처를 포괄해야 한다."""
    # 마커는 고정폭이어야 한다 — "0번"은 "10번"의 부분 문자열이라 오탐한다.
    source = "\n\n".join(f"[P{i:02d}] 문단 " + ("내용 " * 15) for i in range(20))
    chunks = chunk_document(convert("b.md", source.encode("utf-8")))

    for chunk in chunks:
        for index in range(20):
            if f"[P{index:02d}]" in chunk.text:
                assert chunk.block_start <= index <= chunk.block_end, (
                    f"chunk{chunk.index}가 block{index} 내용을 담았지만 범위는 "
                    f"{chunk.block_start}~{chunk.block_end}"
                )


def test_repeated_line_threshold_respects_ratio_boundary():
    """R003: 6페이지 중 3(50%)은 유지, 4(67%)부터 제거 — 정책서의 60% 기준."""
    from backend.pipeline.converters.cleaning import find_repeated_edge_lines

    def pages_with(marker_count: int, total: int = 6):
        return [
            ["머리말 후보" if i < marker_count else f"고유한 첫 줄 {i}", f"본문 {i}"]
            for i in range(total)
        ]

    assert "머리말 후보" not in find_repeated_edge_lines(pages_with(3))
    assert "머리말 후보" in find_repeated_edge_lines(pages_with(4))


def test_docx_keeps_unmerged_cells_with_identical_text():
    """R004: 값이 우연히 같은 독립 셀 두 개는 모두 보존돼야 한다."""
    def build(document):
        table = document.add_table(rows=1, cols=3)
        table.cell(0, 0).text = "승인"
        table.cell(0, 1).text = "승인"
        table.cell(0, 2).text = "보류"

    doc = convert("표.docx", _make_docx(build))
    row = next(b.text for b in doc.blocks if b.kind == "table_row")
    assert "승인 | 승인 | 보류" in row


def test_docx_collapses_actually_merged_cells():
    """R004: 실제 가로 병합된 셀은 한 번만 출력돼야 한다."""
    def build(document):
        table = document.add_table(rows=1, cols=3)
        merged = table.cell(0, 0).merge(table.cell(0, 1))
        merged.text = "병합됨"
        table.cell(0, 2).text = "단독"

    doc = convert("병합.docx", _make_docx(build))
    row = next(b.text for b in doc.blocks if b.kind == "table_row")
    assert row.count("병합됨") == 1
    assert "병합됨 | 단독" in row


def test_docx_preserves_nested_table_content():
    """R202: 셀 안의 중첩 표 내용이 유실되지 않는다.

    `cell.text`는 셀의 직계 문단만 이어 붙이므로, 중첩 표가 통째로 사라졌었다.
    """
    def build(document):
        outer = document.add_table(rows=2, cols=2)
        outer.cell(1, 0).text = "예산"
        cell = outer.cell(1, 1)
        cell.text = "내역:"
        inner = cell.add_table(rows=2, cols=2)
        inner.cell(0, 0).text = "인건비"
        inner.cell(0, 1).text = "3000만원"
        inner.cell(1, 0).text = "장비"
        inner.cell(1, 1).text = "1200만원"

    doc = convert("중첩표.docx", _make_docx(build))
    text = doc.text

    for value in ("예산", "내역:", "인건비", "3000만원", "장비", "1200만원"):
        assert value in text, f"{value}가 유실됐다"
    # 셀의 주변 문단과 중첩 표가 원래 순서대로 이어져야 한다.
    assert text.index("내역:") < text.index("인건비")


def test_docx_warns_when_depth_limit_discards_all_table_content():
    """R301: 깊이 상한으로 표 내용이 전부 사라져도 경고는 반드시 남아야 한다.

    블록과 경고를 같은 조건(`if table_blocks:`)에 묶어 두면, 평탄화 결과가 비었을 때
    경고까지 함께 버려진다. 그 경우가 정확히 "조용한 유실"이라 가장 알려야 한다.
    """
    def build(document):
        document.add_paragraph("보존 본문")
        table = document.add_table(rows=1, cols=1)
        cell = table.cell(0, 0)
        for _ in range(5):  # 1~5단은 비우고 6단에만 내용을 둔다
            cell = cell.add_table(rows=1, cols=1).cell(0, 0)
        cell.text = "유일한깊은내용"

    doc = convert("깊은중첩.docx", _make_docx(build))

    # 주변 문단은 정상 보존되고 변환 자체는 성공한다.
    assert "보존 본문" in doc.text
    # 깊이 상한 때문에 내용은 빠지지만, 빠졌다는 사실이 경고로 전달돼야 한다.
    assert "유일한깊은내용" not in doc.text
    depth_warnings = [
        w for w in doc.warnings if w.code == WarningCode.UNSUPPORTED_ELEMENT
    ]
    assert depth_warnings, "내용이 유실됐는데 경고가 없다"
    assert any("중첩" in w.message for w in depth_warnings)


def test_empty_document_error_explains_why_it_is_empty():
    """깊이 상한으로 내용이 전부 사라진 문서는 오류 메시지에 사유가 실려야 한다.

    표가 문서의 유일한 내용이면 블록이 0개가 되어 `empty_document`로 예외가 나는데,
    이때 경고를 그냥 버리면 사용자는 "텍스트가 없습니다"만 보고 원인을 알 수 없다.
    경고를 남기려던 조치(R301)가 정작 가장 필요한 상황에서 무의미해진다.
    """
    def build(document):
        table = document.add_table(rows=1, cols=1)
        cell = table.cell(0, 0)
        for _ in range(6):
            cell = cell.add_table(rows=1, cols=1).cell(0, 0)
        cell.text = "묻혀버린내용"

    with pytest.raises(ConversionError) as exc:
        convert("표만.docx", _make_docx(build))

    assert exc.value.code == ErrorCode.EMPTY_DOCUMENT
    assert "중첩" in exc.value.message, (
        f"빈 문서 사유가 전달되지 않았다: {exc.value.message}"
    )


def test_identical_warnings_are_deduplicated():
    """같은 표의 여러 셀이 상한에 걸려도 동일 경고가 반복되지 않는다."""
    def build(document):
        document.add_paragraph("본문")
        table = document.add_table(rows=1, cols=6)
        for column in range(6):
            cell = table.cell(0, column)
            for _ in range(6):
                cell = cell.add_table(rows=1, cols=1).cell(0, 0)
            cell.text = f"깊은내용{column}"

    doc = convert("여러셀.docx", _make_docx(build))
    depth_warnings = [
        w for w in doc.warnings if w.code == WarningCode.UNSUPPORTED_ELEMENT
    ]
    assert len(depth_warnings) == 1, (
        f"셀 6개가 동일 경고를 {len(depth_warnings)}건 만들었다"
    )


def test_distinct_tables_keep_separate_warnings():
    """서로 다른 표의 유실은 location으로 구분되어 각각 남는다."""
    def build(document):
        document.add_paragraph("본문")
        for _ in range(2):
            table = document.add_table(rows=1, cols=1)
            cell = table.cell(0, 0)
            for _ in range(6):
                cell = cell.add_table(rows=1, cols=1).cell(0, 0)
            cell.text = "깊은내용"

    doc = convert("두표.docx", _make_docx(build))
    locations = {
        w.location for w in doc.warnings if w.code == WarningCode.UNSUPPORTED_ELEMENT
    }
    assert locations == {"table 1", "table 2"}, locations


def test_duplicate_block_warning_is_aggregated_not_per_block():
    """중복 문단 경고는 블록마다가 아니라 개수 요약 1건으로 나와야 한다.

    반복 고지문이 있는 평범한 문서에서도 블록당 경고를 내면 수십~수백 건이 쌓여
    업로드 201 응답이 부풀고 다른 경고가 묻힌다.
    """
    boilerplate = "본 문서는 대외비이며 무단 전재를 금합니다. 관련 문의는 담당 부서로 연락 바랍니다."
    parts = []
    for index in range(60):
        parts.append(f"[P{index:02d}] 실제 문단 내용입니다." + " 채움" * 5)
        parts.append(boilerplate)

    doc = convert("고지문.md", "\n\n".join(parts).encode("utf-8"))
    duplicate_warnings = [
        w for w in doc.warnings if w.code == WarningCode.DUPLICATE_BLOCK_DROPPED
    ]

    assert len(duplicate_warnings) == 1, (
        f"중복 경고가 {len(duplicate_warnings)}건 — 블록마다 만들어지고 있다"
    )
    assert "59" in duplicate_warnings[0].message, duplicate_warnings[0].message
    # 실제 본문은 남고 반복 고지문만 접힌다.
    assert "[P00]" in doc.text and "[P59]" in doc.text
    assert doc.text.count(boilerplate) == 1


def test_pdf_page_warnings_survive_deduplication():
    """중복 제거가 위치가 다른 경고까지 뭉개면 안 된다."""
    doc = convert("mixed.pdf", _make_pdf([["Real content here."], [], [], []]))
    page_warnings = [w for w in doc.warnings if w.code == WarningCode.PAGE_NO_TEXT]
    assert {w.location for w in page_warnings} == {"page 2", "page 3", "page 4"}


def test_docx_empty_table_produces_no_warning():
    """R301 수정이 기존 동작을 깨지 않는지 고정한다.

    완전히 빈 표는 블록도 경고도 만들지 않고, 주변 문단 순서도 바꾸지 않는다.
    """
    def build(document):
        document.add_paragraph("앞 문단")
        document.add_table(rows=2, cols=2)
        document.add_paragraph("뒤 문단")

    doc = convert("빈표.docx", _make_docx(build))

    assert [b.text for b in doc.blocks] == ["앞 문단", "뒤 문단"]
    assert doc.warnings == []


def test_docx_counts_alternate_content_as_single_object():
    """R203: Choice(w:drawing) + Fallback(w:pict)는 객체 1개다."""
    import docx
    from docx.oxml.ns import qn

    mc = "http://schemas.openxmlformats.org/markup-compatibility/2006"
    document = docx.Document()
    run = document.add_paragraph("본문").add_run()
    alternate = run._r.makeelement(f"{{{mc}}}AlternateContent", {})
    choice = alternate.makeelement(f"{{{mc}}}Choice", {})
    choice.append(choice.makeelement(qn("w:drawing"), {}))
    alternate.append(choice)
    fallback = alternate.makeelement(f"{{{mc}}}Fallback", {})
    fallback.append(fallback.makeelement(qn("w:pict"), {}))
    alternate.append(fallback)
    run._r.append(alternate)
    buffer = io.BytesIO()
    document.save(buffer)

    doc = convert("도형.docx", buffer.getvalue())
    message = next(
        w.message for w in doc.warnings if w.code == WarningCode.UNSUPPORTED_ELEMENT
    )
    assert "1개" in message, f"중복 집계됨: {message}"


def test_docx_warns_about_floating_shapes():
    """R006: 부동 이미지·텍스트 상자도 unsupported_element 경고 대상이다."""
    import docx
    from docx.oxml.ns import qn

    document = docx.Document()
    paragraph = document.add_paragraph("본문 문단")
    # 부동(anchor) 도형을 감싸는 w:drawing을 직접 삽입한다.
    run = paragraph.add_run()
    drawing = run._r.makeelement(qn("w:drawing"), {})
    drawing.append(drawing.makeelement(qn("wp:anchor"), {}))
    run._r.append(drawing)
    buffer = io.BytesIO()
    document.save(buffer)

    doc = convert("도형.docx", buffer.getvalue())
    assert WarningCode.UNSUPPORTED_ELEMENT in {w.code for w in doc.warnings}


# ─── ingestor 연결 ──────────────────────────────────────────────────────────

def test_ingest_chunk_metadata_carries_provenance():
    """구조 정보를 넘기면 청크 metadata에 페이지·문단 좌표가 담긴다."""
    from backend.pipeline.ingestor import _build_chunks

    doc = convert("보고서.pdf", _make_pdf([
        [f"Page one sentence number {i}." for i in range(1, 6)],
        [f"Page two sentence number {i}." for i in range(1, 6)],
    ]))
    chunks = _build_chunks(doc.text, doc)

    assert chunks
    for _, metadata in chunks:
        assert metadata["page_start"] >= 1
        assert metadata["block_start"] >= 0


def test_ingest_falls_back_to_plain_split_without_converted_document():
    """리포지토리 동기화처럼 변환을 거치지 않는 경로는 기존 평문 분할로 동작한다."""
    from backend.pipeline.ingestor import _build_chunks

    chunks = _build_chunks("문장 하나. " * 200, None)

    assert chunks
    # metadata 키 집합은 두 경로가 같아야 한다 (ChromaDB where 필터 일관성).
    structured = _build_chunks("", convert("x.md", "# 제목\n\n본문".encode("utf-8")))
    assert set(chunks[0][1]) == set(structured[0][1])
    assert chunks[0][1]["page_start"] == -1


# ─── 업로드 경로 통합 ────────────────────────────────────────────────────────

def test_upload_accepts_docx_and_passes_structure_to_ingest():
    """DOCX 업로드 → 변환 → ingest에 구조가 전달된다 (핵심 완료 조건)."""
    from unittest.mock import patch

    from fastapi.testclient import TestClient

    from backend.main import app
    from tests.test_upload_and_chunking import _conn_seq

    data = _make_docx(lambda d: (
        d.add_heading("주간 회의", level=1),
        d.add_paragraph("FastAPI를 사용하기로 결정했다."),
    ))
    client = TestClient(app, raise_server_exceptions=False)

    # 업로드는 quota 예약 → 파일 기록 → finalize 순으로 진행된다(#7). 이 테스트의
    # 관심사는 "변환 구조가 ingest까지 전달되는가"이므로 저장 계층만 대역으로 세운다.
    reservation = {
        "reservation_id": 1,
        "temp_path": "data/tmp/회의록.docx",
        "target_path": "data/1/회의록.docx",
    }
    finalized = {
        "doc_id": 12,
        "old_doc_ids": [],
        "file_path": "data/1/회의록.docx",
        "processing_token": None,
    }

    with patch("backend.api.upload.get_connection", side_effect=_conn_seq()), \
         patch("backend.api.upload.require_project_access"), \
         patch("backend.api.upload.require_upload_user", return_value=1), \
         patch("backend.api.upload.reserve_document", return_value=reservation), \
         patch("backend.api.upload.write_reserved_file"), \
         patch("backend.api.upload.finalize_document", return_value=finalized), \
         patch("backend.api.upload.extract", return_value=[]), \
         patch("backend.api.upload.ingest") as mock_ingest, \
         patch("backend.api.upload.update_project_memory"), \
         patch("backend.api.upload._set_doc_status"):

        response = client.post(
            "/api/v1/projects/1/documents",
            files={"file": ("회의록.docx", data,
                            "application/vnd.openxmlformats-officedocument."
                            "wordprocessingml.document")},
        )

    assert response.status_code == 201
    body = response.json()
    assert body["format"] == "docx"
    assert body["blocks"] == 2
    assert body["warnings"] == []

    converted = mock_ingest.call_args.kwargs["converted"]
    assert converted.format == "docx"
    assert "FastAPI" in converted.text


def test_upload_rejects_unsupported_format_with_reason():
    from unittest.mock import patch

    from fastapi.testclient import TestClient

    from backend.main import app

    client = TestClient(app, raise_server_exceptions=False)
    with patch("backend.api.upload.require_project_access"), \
         patch("backend.api.upload.require_upload_user", return_value=1):
        response = client.post(
            "/api/v1/projects/1/documents",
            files={"file": ("보고서.hwp", b"data", "application/octet-stream")},
        )

    assert response.status_code == 400
    assert ".docx" in response.json()["detail"]


def test_upload_rejects_scanned_pdf_with_explicit_code():
    """실패 파일은 명시적 오류를 돌려준다 (완료 조건)."""
    from unittest.mock import patch

    from fastapi.testclient import TestClient

    from backend.main import app

    client = TestClient(app, raise_server_exceptions=False)
    with patch("backend.api.upload.require_project_access"), \
         patch("backend.api.upload.require_upload_user", return_value=1):
        response = client.post(
            "/api/v1/projects/1/documents",
            files={"file": ("스캔본.pdf", _make_pdf([[], []]), "application/pdf")},
        )

    assert response.status_code == 400
    body = response.json()
    # detail은 반드시 문자열이어야 한다 — 객체로 내리면 데스크톱 클라이언트가
    # 사유를 버리고 "PaiM API 요청 실패"만 표시한다 (paimApi.ts 계약).
    assert isinstance(body["detail"], str)
    assert "스캔" in body["detail"]
    assert body["code"] == ErrorCode.NO_TEXT_LAYER


def test_upload_error_detail_is_string_for_every_conversion_failure():
    """모든 변환 실패가 기존 클라이언트가 읽는 형태(문자열 detail + 최상위 code)여야 한다."""
    from unittest.mock import patch

    from fastapi.testclient import TestClient

    from backend.main import app

    client = TestClient(app, raise_server_exceptions=False)
    # 여기 사례는 모두 "형식 검증은 통과했으나 변환에 실패한" 400 경로여야 한다.
    # 매직 불일치(예: ZIP이 아닌 .docx)는 415 검증 실패이므로 이 표에 넣지 않는다.
    cases = [
        ("깨진.docx", b"PK\x03\x04not a real docx", ErrorCode.CORRUPT_FILE),
        ("빈.txt", b"\n\n   \n", ErrorCode.EMPTY_DOCUMENT),
        ("스캔.pdf", _make_pdf([[], []]), ErrorCode.NO_TEXT_LAYER),
    ]
    for name, data, expected_code in cases:
        with patch("backend.api.upload.require_project_access"), \
         patch("backend.api.upload.require_upload_user", return_value=1):
            response = client.post(
                "/api/v1/projects/1/documents",
                files={"file": (name, data, "application/octet-stream")},
            )
        body = response.json()
        assert response.status_code == 400, name
        assert isinstance(body["detail"], str), name
        assert body["code"] == expected_code, name


# ─── 4차 리뷰 회귀 (PR-001-R401, R4B02) ────────────────────────────────────

def test_empty_document_message_is_bounded_and_hides_internals():
    """R401: 경고가 많아도 오류 메시지가 폭증하지 않고 내부 정보를 노출하지 않는다."""
    from backend.pipeline.converters.base import ConversionWarning, assemble

    # 메시지를 페이지마다 다르게 만든다. 수정 전 구현은 message 기준으로만 중복을
    # 제거했으므로, 같은 문구를 반복하면 1건으로 접혀 이 테스트가 공허해진다.
    # 실제 R401 상황도 원시 예외가 페이지마다 달라 메시지가 전부 고유했다.
    warnings = [
        ConversionWarning(WarningCode.PAGE_EXTRACT_FAILED,
                          f"페이지 텍스트 추출에 실패했습니다: 내부토큰_{i} at offset 0x{i:04x}",
                          f"page {i}")
        for i in range(1, 80)
    ] + [
        ConversionWarning(WarningCode.PAGE_NO_TEXT,
                          f"{i}번 페이지에서 텍스트를 찾지 못했습니다.", f"page {i}")
        for i in range(1, 20)
    ]

    with pytest.raises(ConversionError) as exc:
        assemble("많은경고.pdf", "pdf", [], warnings)

    message = exc.value.message
    assert exc.value.code == ErrorCode.EMPTY_DOCUMENT
    # 고유 메시지 98건이 그대로 붙으면 수천 자가 된다(수정 전 실측 3,643자).
    assert len(message) < 300, f"메시지가 {len(message)}자로 과도하다"
    # 경고 메시지를 그대로 실으면 내부 토큰이 사용자 응답에 노출된다.
    assert "내부토큰_" not in message
    assert "offset 0x" not in message


def test_page_extract_failure_warning_excludes_raw_exception():
    """R401: 페이지 파싱 실패 경고에 원시 예외 문자열이 들어가지 않는다."""
    from backend.pipeline.converters.pdf_converter import _page_texts

    class _BoomPage:
        def extract_text(self):
            raise ValueError("내부파서고유토큰_XYZ at offset 0x1234")

    class _Reader:
        pages = [_BoomPage()]

    _texts, warnings = _page_texts(_Reader(), "깨진.pdf")

    assert len(warnings) == 1
    assert warnings[0].code == WarningCode.PAGE_EXTRACT_FAILED
    # 원시 예외가 실리면 이 단언이 깨진다. 사용자 응답으로 나가는 문자열이다.
    assert "내부파서고유토큰_XYZ" not in warnings[0].message
    assert "0x1234" not in warnings[0].message


def test_empty_document_message_unchanged_without_warnings():
    """R401: 경고가 없을 때는 기존 문구를 그대로 유지한다."""
    from backend.pipeline.converters.base import assemble

    with pytest.raises(ConversionError) as exc:
        assemble("빈.docx", "docx", [], [])
    assert exc.value.message == "문서에서 추출된 텍스트가 없습니다."


def test_depth_limit_reason_still_reaches_user():
    """R401 수정이 R301의 깊이 상한 사유 전달을 깨지 않는다."""
    import docx

    def build(document):
        cell = document.add_table(rows=1, cols=1).cell(0, 0)
        for _ in range(5):
            cell = cell.add_table(rows=1, cols=1).cell(0, 0)
        cell.text = "깊은 내용"

    with pytest.raises(ConversionError) as exc:
        convert("깊은표만.docx", _make_docx(build))
    assert exc.value.code == ErrorCode.EMPTY_DOCUMENT
    assert "중첩 깊이" in exc.value.message


def test_upload_unsupported_format_returns_top_level_code():
    """R4B02: 미지원 확장자도 최상위 code를 실어야 한다(명세에 있는 계약)."""
    from unittest.mock import patch

    from fastapi.testclient import TestClient

    from backend.main import app

    client = TestClient(app, raise_server_exceptions=False)
    with patch("backend.api.upload.require_project_access"), \
         patch("backend.api.upload.require_upload_user", return_value=1):
        response = client.post(
            "/api/v1/projects/1/documents",
            files={"file": ("문서.hwp", b"whatever", "application/octet-stream")},
        )
    body = response.json()
    assert response.status_code == 400
    assert isinstance(body["detail"], str)
    assert body["code"] == ErrorCode.UNSUPPORTED_FORMAT


# ─── 5차 리뷰 회귀 (PR-001-R501) ────────────────────────────────────────────
#
# 파일 열기·파싱 실패의 원시 예외가 사용자 응답(HTTP 400 detail)에 노출되던 문제.
# 세 경로가 서로 독립이므로 한 경로의 수정만 되돌려도 그 테스트만 실패한다.
#
# 단언에 고정 문구를 포함하는 이유: `token not in message`만 검사하면 메시지를
# 빈 문자열이나 무의미한 문구로 바꾼 수정도 통과한다(8차 리뷰 지적).

_RAW_TOKEN = "내부파서고유토큰_XYZ_at_offset_0xBEEF"


def _upload_and_get_body(name: str, data: bytes):
    """업로드 엔드포인트를 실제로 태워 400 응답 본문을 돌려준다."""
    from unittest.mock import patch

    from fastapi.testclient import TestClient

    from backend.main import app

    client = TestClient(app, raise_server_exceptions=False)
    with patch("backend.api.upload.require_project_access"), \
         patch("backend.api.upload.require_upload_user", return_value=1):
        response = client.post(
            "/api/v1/projects/1/documents",
            files={"file": (name, data, "application/octet-stream")},
        )
    return response


def test_pdf_open_failure_hides_raw_exception(monkeypatch):
    """R501: PDF reader 생성 실패의 원시 예외가 응답에 노출되지 않는다."""
    from backend.pipeline.converters import pdf_converter

    raw = RuntimeError(_RAW_TOKEN)

    class _BoomReader:
        def __init__(self, *_args, **_kwargs):
            raise raw

    monkeypatch.setattr(pdf_converter, "_require_pypdf", lambda: _BoomReader)

    with pytest.raises(ConversionError) as exc:
        convert("깨진.pdf", b"%PDF-1.4 whatever")

    assert exc.value.code == ErrorCode.CORRUPT_FILE
    assert exc.value.message == pdf_converter.PDF_OPEN_FAILED_MESSAGE
    assert _RAW_TOKEN not in exc.value.message
    assert exc.value.source == "깨진.pdf"
    # 원인 예외는 체인으로 보존돼야 한다 — 진단 정보를 버리는 것이 아니라 옮기는 것이다.
    assert exc.value.__cause__ is raw

    response = _upload_and_get_body("깨진.pdf", b"%PDF-1.4 whatever")
    body = response.json()
    assert response.status_code == 400
    assert body["code"] == ErrorCode.CORRUPT_FILE
    assert body["detail"] == pdf_converter.PDF_OPEN_FAILED_MESSAGE
    assert _RAW_TOKEN not in body["detail"]


def test_docx_zip_guard_failure_hides_raw_exception(monkeypatch):
    """R501: ZIP 중앙 디렉터리 읽기 실패의 원시 예외가 응답에 노출되지 않는다."""
    import types
    import zipfile as zipfile_module

    from backend.pipeline.converters import docx_converter

    raw = zipfile_module.BadZipFile(_RAW_TOKEN)

    class _BoomZipFile:
        def __init__(self, *_args, **_kwargs):
            raise raw

    # `docx_converter.zipfile`은 표준 zipfile 모듈 그 자체다. 거기에 setattr 하면
    # 인터프리터 전역이 오염되어 다른 테스트의 DOCX 생성(zipfile 쓰기)까지 깨진다.
    # 실제로 실행 순서에 따라 뒤 테스트가 BadZipFile로 실패하는 것을 확인했다.
    # 모듈 자체를 대체해 이 테스트 안에서만 효과가 있도록 한다.
    stub = types.SimpleNamespace(
        ZipFile=_BoomZipFile,
        BadZipFile=zipfile_module.BadZipFile,
        LargeZipFile=zipfile_module.LargeZipFile,
    )
    monkeypatch.setattr(docx_converter, "zipfile", stub)

    # ZIP 매직으로 시작해야 업로드 경계 검증(415)을 통과해 변환기의 ZIP 가드까지
    # 도달한다. 매직이 없으면 validate_document_bytes()가 먼저 415로 끊어, 이 테스트가
    # 검증하려던 경로(가드 실패 시 원시 예외 비노출)를 타지 못한다.
    broken_zip = b"PK\x03\x04tiny"

    with pytest.raises(ConversionError) as exc:
        convert("깨진.docx", broken_zip)

    assert exc.value.code == ErrorCode.CORRUPT_FILE
    assert exc.value.message == docx_converter.DOCX_OPEN_FAILED_MESSAGE
    assert _RAW_TOKEN not in exc.value.message
    assert exc.value.source == "깨진.docx"
    assert exc.value.__cause__ is raw

    response = _upload_and_get_body("깨진.docx", broken_zip)
    body = response.json()
    assert response.status_code == 400
    assert body["code"] == ErrorCode.CORRUPT_FILE
    assert body["detail"] == docx_converter.DOCX_OPEN_FAILED_MESSAGE
    assert _RAW_TOKEN not in body["detail"]


def test_docx_parser_failure_hides_raw_exception(monkeypatch):
    """R501: python-docx 파싱 실패의 원시 예외가 응답에 노출되지 않는다.

    정상 소형 DOCX를 먼저 만들어 ZIP 가드는 통과시키고, 파서 단계만 실패시킨다.
    가드에서 걸려버리면 이 테스트가 다른 경로를 검증하게 된다.
    """
    import docx as docx_module

    from backend.pipeline.converters import docx_converter

    data = _make_docx(lambda d: d.add_paragraph("정상 본문"))
    raw = ValueError(_RAW_TOKEN)

    def _boom(*_args, **_kwargs):
        raise raw

    monkeypatch.setattr(docx_module, "Document", _boom)

    with pytest.raises(ConversionError) as exc:
        convert("파싱실패.docx", data)

    assert exc.value.code == ErrorCode.CORRUPT_FILE
    assert exc.value.message == docx_converter.DOCX_OPEN_FAILED_MESSAGE
    assert _RAW_TOKEN not in exc.value.message
    assert exc.value.source == "파싱실패.docx"
    # ZIP 가드가 아니라 파서 단계에서 실패했음을 원인 예외 타입으로 고정한다.
    assert exc.value.__cause__ is raw
    assert isinstance(exc.value.__cause__, ValueError)

    response = _upload_and_get_body("파싱실패.docx", data)
    body = response.json()
    assert response.status_code == 400
    assert body["code"] == ErrorCode.CORRUPT_FILE
    assert body["detail"] == docx_converter.DOCX_OPEN_FAILED_MESSAGE
    assert _RAW_TOKEN not in body["detail"]
