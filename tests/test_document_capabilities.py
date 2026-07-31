import base64
import io
from pathlib import Path

import pytest
from fastapi import HTTPException
from fastapi.testclient import TestClient

from backend.api import query as query_api
from backend.document_content import (
    ALLOWED_SUFFIXES,
    DOCUMENT_PARSERS,
    PROJECT_DOCUMENT_MAX_FILE_BYTES,
    QUERY_ATTACHMENT_MAX_FILE_BYTES,
    QUERY_ATTACHMENT_MAX_TOTAL_BYTES,
    extract_document_text,
    supported_extensions,
)
from backend.main import app
from backend.pipeline.converters import supported_suffixes

client = TestClient(app, raise_server_exceptions=False)
EXPECTED_DOCUMENT_SUFFIXES = {".docx", ".markdown", ".md", ".pdf", ".txt"}


def _query_attachment_context(attachments):
    evidence = query_api._prepare_attachment_evidence(attachments)
    return query_api._render_attachment_evidence(evidence)


def _make_docx(text: str) -> bytes:
    import docx

    document = docx.Document()
    document.add_heading("회의 결과", level=1)
    document.add_paragraph(text)
    table = document.add_table(rows=1, cols=2)
    table.cell(0, 0).text = "담당"
    table.cell(0, 1).text = "김개발"
    output = io.BytesIO()
    document.save(output)
    return output.getvalue()


def _make_pdf(text: str) -> bytes:
    stream = f"BT /F1 12 Tf 72 720 Td ({text}) Tj ET"
    objects = [
        "<< /Type /Catalog /Pages 2 0 R >>",
        "<< /Type /Pages /Kids [3 0 R] /Count 1 >>",
        "<< /Type /Page /Parent 2 0 R /MediaBox [0 0 612 792] "
        "/Resources << /Font << /F1 5 0 R >> >> /Contents 4 0 R >>",
        f"<< /Length {len(stream)} >>\nstream\n{stream}\nendstream",
        "<< /Type /Font /Subtype /Type1 /BaseFont /Helvetica >>",
    ]
    output = io.BytesIO()
    output.write(b"%PDF-1.4\n")
    offsets = []
    for number, content in enumerate(objects, 1):
        offsets.append(output.tell())
        output.write(f"{number} 0 obj\n{content}\nendobj\n".encode("ascii"))
    xref = output.tell()
    output.write(f"xref\n0 {len(objects) + 1}\n0000000000 65535 f \n".encode())
    for offset in offsets:
        output.write(f"{offset:010d} 00000 n \n".encode())
    output.write(
        f"trailer\n<< /Size {len(objects) + 1} /Root 1 0 R >>\n"
        f"startxref\n{xref}\n%%EOF\n".encode()
    )
    return output.getvalue()


def test_capabilities_exactly_match_converter_registry():
    response = client.get("/api/v1/capabilities")
    assert response.status_code == 200
    body = response.json()
    converter_suffixes = set(supported_suffixes())
    assert converter_suffixes == EXPECTED_DOCUMENT_SUFFIXES
    assert set(DOCUMENT_PARSERS) == converter_suffixes
    assert ALLOWED_SUFFIXES == converter_suffixes

    expected = sorted(suffix.removeprefix(".") for suffix in converter_suffixes)
    assert supported_extensions() == expected
    assert body == {
        "schema_version": 1,
        "project_documents": {
            "extensions": expected,
            "max_file_bytes": PROJECT_DOCUMENT_MAX_FILE_BYTES,
        },
        "query_attachments": {
            "extensions": expected,
            "max_file_bytes": QUERY_ATTACHMENT_MAX_FILE_BYTES,
            "max_total_bytes": QUERY_ATTACHMENT_MAX_TOTAL_BYTES,
        },
        "desktop_chat": {
            "storage": "local_only",
            "server_persistence": False,
            "legacy_session_api": "deprecated",
        },
    }


def test_capabilities_requires_authentication(monkeypatch):
    monkeypatch.setenv("PAIM_AUTH_MODE", "jwt")
    response = client.get("/api/v1/capabilities")
    assert response.status_code == 401


def test_unsupported_query_attachment_returns_400():
    attachment = query_api.QueryAttachment(
        filename="malware.exe",
        content_base64=base64.b64encode(b"content").decode(),
    )
    with pytest.raises(HTTPException) as exc:
        _query_attachment_context([attachment])
    assert exc.value.status_code == 400


def test_query_attachment_size_boundary(monkeypatch):
    monkeypatch.setattr(query_api, "QUERY_ATTACHMENT_MAX_FILE_BYTES", 3)
    monkeypatch.setattr(query_api, "QUERY_ATTACHMENT_MAX_TOTAL_BYTES", 3)
    exact = query_api.QueryAttachment(
        filename="exact.txt",
        content_base64=base64.b64encode(b"abc").decode(),
    )
    assert "abc" in _query_attachment_context([exact])[0]

    oversized = query_api.QueryAttachment(
        filename="large.txt",
        content_base64=base64.b64encode(b"abcd").decode(),
    )
    with pytest.raises(HTTPException) as exc:
        _query_attachment_context([oversized])
    assert exc.value.status_code == 413


def test_query_attachment_total_decoded_size_returns_413(monkeypatch):
    monkeypatch.setattr(query_api, "QUERY_ATTACHMENT_MAX_FILE_BYTES", 4)
    monkeypatch.setattr(query_api, "QUERY_ATTACHMENT_MAX_TOTAL_BYTES", 5)
    attachments = [
        query_api.QueryAttachment(
            filename=f"{index}.txt",
            content_base64=base64.b64encode(b"abc").decode(),
        )
        for index in range(2)
    ]
    with pytest.raises(HTTPException) as exc:
        _query_attachment_context(attachments)
    assert exc.value.status_code == 413


def test_query_total_size_is_checked_after_text_context_is_full(monkeypatch):
    monkeypatch.setattr(query_api, "QUERY_ATTACHMENT_MAX_FILE_BYTES", 4)
    monkeypatch.setattr(query_api, "QUERY_ATTACHMENT_MAX_TOTAL_BYTES", 5)
    monkeypatch.setattr(query_api, "_ATTACHMENT_MAX_CHARS_TOTAL", 1)
    attachments = [
        query_api.QueryAttachment(
            filename=f"{index}.txt",
            content_base64=base64.b64encode(b"abc").decode(),
        )
        for index in range(2)
    ]
    with pytest.raises(HTTPException) as exc:
        _query_attachment_context(attachments)
    assert exc.value.status_code == 413


def test_advertised_parsers_extract_content():
    assert extract_document_text("notes.md", b"markdown body") == "markdown body"
    assert extract_document_text("notes.markdown", b"long markdown body") == "long markdown body"
    assert extract_document_text("notes.txt", b"text body") == "text body"
    assert "PDF body" in extract_document_text("notes.pdf", _make_pdf("PDF body"))
    docx_text = extract_document_text("meeting.docx", _make_docx("DOCX 본문 추출 성공"))
    assert "DOCX 본문 추출 성공" in docx_text
    assert "담당 | 김개발" in docx_text


def test_frontend_has_no_supported_extension_or_size_fallback():
    source = (Path(__file__).parents[1] / "desktop/src/App.tsx").read_text()
    assert 'extensions: ["md", "txt", "pdf", "docx"]' not in source
    assert "10 * 1024 * 1024" not in source
    assert "getUploadMimeType" not in source

    capability_source = (Path(__file__).parents[1] / "desktop/src/capabilities.ts").read_text()
    assert "extensions.includes(getFileExtension(name))" in capability_source
