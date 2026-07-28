"""DOCX 변환기 (필수 포맷).

python-docx는 문단과 표를 별도 컬렉션으로 노출해서 `document.paragraphs`만 읽으면
표가 통째로 사라지고, 본문 순서도 깨진다. 그래서 body XML을 직접 순회해
문단·표가 원본에 나타난 순서 그대로 Block이 되게 한다.

DOCX에는 페이지 개념이 없다(페이지는 렌더링 시점에 결정된다). 따라서 page는
항상 None이고, 출처 추적은 문단 번호(Block.order)와 제목 경로로 한다.
"""
from __future__ import annotations

import io
import logging
import re
import zipfile

from .base import (
    ConversionError,
    ConversionWarning,
    ErrorCode,
    WarningCode,
    assemble,
)
from .cleaning import drop_duplicate_blocks, is_noise_line, normalize_text

logger = logging.getLogger(__name__)

SUFFIXES = (".docx",)

# 파일 열기·파싱 실패 시 사용자에게 나가는 고정 문구. 원시 예외를 섞지 않는다 —
# 이 문자열이 업로드 400 응답의 detail이 된다. ZIP 중앙 디렉터리 실패와 python-docx
# 파싱 실패는 내부 단계만 다르고 사용자가 취할 조치는 같으므로 문구를 공유한다.
DOCX_OPEN_FAILED_MESSAGE = (
    "DOCX 파일을 열 수 없습니다. 파일이 손상되었거나 올바른 DOCX 파일이 아닐 수 있으니 "
    "확인 후 다시 올려주세요."
)

# 압축 폭탄 방어 한도. DOCX는 ZIP이라 업로드 크기 제한(10MB)이 압축된 바이트에만
# 걸린다. 반복 텍스트만으로도 압축비 250:1이 나와 10MB가 2.5GB로 전개될 수 있고,
# python-docx는 XML 전체를 메모리에 올리므로 워커가 그대로 고갈된다.
# 정상 DOCX의 압축비는 보통 10:1 미만이라 아래 한도는 실사용을 방해하지 않는다.
_MAX_UNCOMPRESSED_BYTES = 200 * 1024 * 1024   # 전체 전개 크기
_MAX_ENTRY_BYTES = 100 * 1024 * 1024          # 엔트리 1개 전개 크기
_MAX_COMPRESSION_RATIO = 120                  # 전체 압축비

# Markup Compatibility 네임스페이스. 도형 하나를 DrawingML(Choice)과 VML(Fallback)로
# 함께 담는 컨테이너라, 비텍스트 요소를 셀 때 논리 객체 1개로 취급해야 한다.
_MC_ALTERNATE_CONTENT = (
    "{http://schemas.openxmlformats.org/markup-compatibility/2006}AlternateContent"
)

# "Heading 1", "제목 1", "개요 1" 등 스타일 이름에서 레벨 숫자를 뽑는다.
_HEADING_STYLE = re.compile(r"^(?:heading|title|제목|개요)\s*(\d+)?$", re.IGNORECASE)
_LIST_STYLE = re.compile(r"(list|목록|bullet|번호)", re.IGNORECASE)


def _require_python_docx():
    try:
        import docx  # noqa: F401
        from docx.oxml.ns import qn
        from docx.table import Table
        from docx.text.paragraph import Paragraph
    except ImportError as exc:
        raise ConversionError(
            ErrorCode.MISSING_DEPENDENCY,
            "DOCX 변환에 필요한 python-docx가 설치되어 있지 않습니다.",
        ) from exc
    return docx, qn, Table, Paragraph


def _iter_body(document, qn, Table, Paragraph):
    """body의 자식 엘리먼트를 원본 순서대로 문단/표 객체로 넘겨준다."""
    for child in document.element.body.iterchildren():
        if child.tag == qn("w:p"):
            yield Paragraph(child, document)
        elif child.tag == qn("w:tbl"):
            yield Table(child, document)


def _style_name_map(document) -> dict[str, str]:
    """styleId → 표시 이름 사전을 문서당 한 번만 만든다.

    `paragraph.style`을 문단마다 읽으면 python-docx가 문단마다 기본 스타일을
    다시 해석해(styles.default_for) 대형 문서에서 변환 시간의 대부분을 차지한다.
    스타일 수는 보통 수십 개뿐이라 미리 한 번 펼쳐 두는 편이 훨씬 싸다.

    styleId를 그대로 쓰지 않고 이름으로 바꾸는 이유는, 국내에서 작성된 DOCX가
    styleId를 `a3` 같은 불투명한 값으로 갖고 이름에만 `제목 1`이 담기는 경우가
    흔하기 때문이다.
    """
    names: dict[str, str] = {}
    try:
        for style in document.styles:
            style_id = getattr(style, "style_id", None)
            if style_id:
                names[style_id] = (getattr(style, "name", None) or style_id).strip()
    except Exception:
        # 스타일 파트가 손상된 문서에서도 본문 변환은 계속돼야 한다.
        logger.debug("DOCX 스타일 목록을 읽지 못했습니다. styleId로 대체합니다.", exc_info=True)
    return names


def _paragraph_kind(paragraph, style_names: dict[str, str]) -> tuple[str, int | None]:
    """문단 스타일에서 (kind, level)을 판정한다.

    pStyle이 없는 문단은 기본 스타일(본문)이므로 이름 해석 없이 paragraph로 본다.
    """
    style_name = ""
    try:
        properties = paragraph._p.pPr
        if properties is not None and properties.pStyle is not None:
            style_id = properties.pStyle.val
            style_name = style_names.get(style_id, style_id or "").strip()
    except Exception:
        # 스타일 참조가 깨진 문서에서도 문단 자체는 살려 본문으로 취급한다.
        style_name = ""

    heading = _HEADING_STYLE.match(style_name)
    if heading:
        level = int(heading.group(1)) if heading.group(1) else 1
        return "heading", level

    # numPr(번호매기기 속성)이 있으면 스타일 이름과 무관하게 목록으로 본다.
    try:
        if paragraph._p.pPr is not None and paragraph._p.pPr.numPr is not None:
            return "list_item", None
    except AttributeError:
        pass
    if _LIST_STYLE.search(style_name):
        return "list_item", None
    return "paragraph", None


# 중첩 표 재귀 상한. 실제 문서에서 3단을 넘는 중첩은 사실상 없고, 악의적으로 깊게
# 중첩한 문서가 스택을 소모하는 것을 막는다.
_MAX_TABLE_DEPTH = 5
# 중첩 표의 행 구분자. 바깥 셀 구분자(" | ")와 시각적으로 구분되어야 읽힌다.
_NESTED_ROW_SEPARATOR = " ; "


def _iter_cell_rows(row, qn, Table, Paragraph):
    """행의 셀들을 병합 접기까지 적용해 돌려준다.

    가로 병합된 셀은 python-docx가 같은 <w:tc> 엘리먼트를 반복해서 돌려준다.
    텍스트가 같다고 접으면 우연히 값이 같은 독립 셀("승인 | 승인")까지 하나로
    뭉개져 열 위치가 밀리므로, 반드시 엘리먼트 동일성으로 판정한다.
    """
    previous_tc = None
    for cell in row.cells:
        if cell._tc is previous_tc:
            continue
        previous_tc = cell._tc
        yield cell


def _cell_text(cell, qn, Table, Paragraph, depth: int, warnings: list, location: str) -> str:
    """셀 내부를 XML 순서대로 읽어 문단과 중첩 표를 모두 텍스트로 만든다.

    `cell.text`는 셀의 **직계 문단만** 이어 붙이므로 중첩 표가 통째로 사라진다.
    자식 엘리먼트를 직접 순회해 w:p와 w:tbl을 원래 순서대로 처리한다.
    """
    parts: list[str] = []
    for child in cell._tc.iterchildren():
        if child.tag == qn("w:p"):
            text = normalize_text(Paragraph(child, cell).text).replace("\n", " ").strip()
            if text:
                parts.append(text)
        elif child.tag == qn("w:tbl"):
            if depth >= _MAX_TABLE_DEPTH:
                # 상한 때문에 버릴 때는 반드시 알린다. 조용히 버리면 이 결함을 되풀이한다.
                # location을 붙여 어느 표인지 짚어주고, 같은 표 안에서 여러 셀이
                # 상한에 걸려도 동일 경고가 반복되지 않게 한다(assemble에서 중복 제거).
                warnings.append(ConversionWarning(
                    WarningCode.UNSUPPORTED_ELEMENT,
                    f"중첩 깊이 {_MAX_TABLE_DEPTH}단을 초과한 표의 내용은 "
                    "변환되지 않았습니다.",
                    location=location,
                ))
                continue
            nested_rows = _flatten_table_rows(
                Table(child, cell), qn, Table, Paragraph, depth + 1, warnings, location
            )
            if nested_rows:
                parts.append(_NESTED_ROW_SEPARATOR.join(nested_rows))
    return " ".join(parts).strip()


def _flatten_table_rows(
    table, qn, Table, Paragraph, depth: int, warnings: list, location: str,
) -> list[str]:
    """표를 행 문자열 목록으로 평탄화한다(중첩 표 포함)."""
    rows: list[str] = []
    for row in table.rows:
        cells = [
            _cell_text(cell, qn, Table, Paragraph, depth, warnings, location)
            for cell in _iter_cell_rows(row, qn, Table, Paragraph)
        ]
        text = " | ".join(c for c in cells if c)
        if text:
            rows.append(text)
    return rows


def _table_blocks(
    table, table_index: int, qn, Table, Paragraph,
) -> tuple[list[dict], ConversionWarning, list[ConversionWarning]]:
    """표를 행 단위 Block으로 평탄화한다.

    표 구조(행/열)를 그대로 담을 자리가 Block 계약에 없으므로 셀을 " | "로 잇는다.

    반환값을 세 개로 나누는 이유는 두 경고의 성격이 다르기 때문이다.
    - `table_flattened`(열 구조 유실 안내)는 출력 블록이 있을 때만 의미가 있다.
    - **내용 유실 경고는 블록이 하나도 안 나와도 반드시 전달돼야 한다.** 두 경고를
      한 목록에 섞어 두면 호출자가 "블록이 있을 때만" 통째로 추가하게 되고,
      정작 내용이 전부 사라진 경우에 경고까지 사라진다(R301).
    """
    location = f"table {table_index + 1}"
    content_loss: list[ConversionWarning] = []
    rows = _flatten_table_rows(
        table, qn, Table, Paragraph, depth=1, warnings=content_loss, location=location,
    )
    blocks = [
        {"kind": "table_row", "text": f"[표{table_index + 1}] {row}"}
        for row in rows
    ]
    flatten_warning = ConversionWarning(
        WarningCode.TABLE_FLATTENED,
        "표를 행 단위 텍스트로 변환했습니다. 열 구조는 보존되지 않습니다.",
        location=location,
    )
    return blocks, flatten_warning, content_loss


def _guard_archive_size(data: bytes, filename: str) -> None:
    """DOCX(ZIP)를 열기 전에 전개 크기와 압축비를 검사한다.

    python-docx에 넘기고 나면 이미 늦다 — 파싱이 시작되는 순간 메모리가 잡힌다.
    ZIP 중앙 디렉터리의 크기 정보만 읽어 그 전에 거절한다.
    """
    try:
        with zipfile.ZipFile(io.BytesIO(data)) as archive:
            entries = archive.infolist()
    except (
        zipfile.BadZipFile,
        zipfile.LargeZipFile,
        # 중앙 디렉터리의 extraction version이 지원 상한을 넘으면 NotImplementedError,
        # UTF-8 플래그와 맞지 않는 파일명이면 UnicodeError가 난다. 가드를 넣기 전에는
        # docx.Document() 주변의 넓은 except가 흡수하던 입력이라, 여기서 정규화하지
        # 않으면 400이어야 할 응답이 500으로 누출된다.
        NotImplementedError,
        UnicodeError,
        ValueError,
        OSError,
    ) as exc:
        # 원시 예외 문자열을 메시지에 넣지 않는다. 이 message는 업로드 400 detail로
        # 그대로 나가므로 ZIP 파서 내부 정보가 사용자에게 노출된다.
        # 진단 정보는 로그와 예외 체인(__cause__)에만 남긴다.
        logger.warning("DOCX ZIP 중앙 디렉터리 읽기 실패 filename=%s", filename, exc_info=True)
        raise ConversionError(
            ErrorCode.CORRUPT_FILE,
            DOCX_OPEN_FAILED_MESSAGE,
            source=filename,
        ) from exc
    # MemoryError 등 자원 고갈 예외는 잡지 않는다 — 입력이 잘못된 게 아니라 서버가
    # 한계에 부딪힌 상황이므로, 정상 입력 오류로 숨기면 장애를 못 보게 된다.

    total = 0
    for entry in entries:
        if entry.file_size > _MAX_ENTRY_BYTES:
            raise ConversionError(
                ErrorCode.FILE_TOO_LARGE,
                "DOCX 내부 항목의 크기가 안전한 처리 한도를 초과했습니다. "
                "문서를 나누거나 내용을 줄여서 다시 올려주세요.",
                source=filename,
            )
        total += entry.file_size

    if total > _MAX_UNCOMPRESSED_BYTES or (
        data and total / len(data) > _MAX_COMPRESSION_RATIO
    ):
        raise ConversionError(
            ErrorCode.FILE_TOO_LARGE,
            "DOCX의 압축 해제 크기가 안전한 처리 한도를 초과했습니다. "
            "문서를 나누거나 내용을 줄여서 다시 올려주세요.",
            source=filename,
        )


def _unsupported_element_count(document, qn) -> int:
    """텍스트로 변환되지 않는 요소 개수를 센다.

    `document.inline_shapes`는 본문 흐름에 박힌 인라인 그림만 센다. 회의록·정책문서에
    흔한 부동 이미지(`wp:anchor`)와 텍스트 상자(`w:txbxContent`)는 거기 잡히지 않아,
    내용이 사라져도 경고가 나가지 않았다. body XML을 직접 세어 누락을 막는다.
    """
    # w:drawing(DrawingML)과 w:pict(구형 VML)는 객체 하나당 하나씩 나타나는 최상위
    # 컨테이너다. 내부의 wp:anchor·w:txbxContent까지 세면 같은 객체를 중복 집계한다.
    #
    # 다만 도형 하나가 mc:AlternateContent 안에 Choice(w:drawing) + Fallback(w:pict)로
    # 함께 저장되는 경우가 있어, 단순 합산하면 1개 객체가 2개로 잡힌다.
    # AlternateContent는 그 자체를 논리 객체 1개로 세고, 그 바깥의 drawing/pict만 더한다.
    total = 0
    try:
        body = document.element.body
        # python-docx의 qn()은 mc 접두어를 모른다(KeyError). Clark 표기를 직접 쓴다.
        alternates = body.findall(".//" + _MC_ALTERNATE_CONTENT)
        total += len(alternates)
        alternate_set = set(alternates)

        for tag in ("w:drawing", "w:pict"):
            try:
                found = body.findall(".//" + qn(tag))
            except (KeyError, ValueError):
                continue
            for element in found:
                if not _has_ancestor(element, alternate_set):
                    total += 1
    except Exception:
        logger.debug("DOCX 비텍스트 요소 집계 실패", exc_info=True)
        return 0
    return total


def _has_ancestor(element, ancestors: set) -> bool:
    """element가 주어진 조상 집합 안에 들어 있는지 판정한다.

    lxml 프록시 객체는 수명이 짧아 id() 비교나 집합 연산을 신뢰하기 어렵다.
    부모를 직접 거슬러 올라가며 확인한다.
    """
    if not ancestors:
        return False
    parent = element.getparent()
    while parent is not None:
        for candidate in ancestors:
            if parent is candidate:
                return True
        parent = parent.getparent()
    return False


def convert(filename: str, data: bytes):
    """DOCX 바이트를 ConvertedDocument로 변환한다."""
    docx, qn, Table, Paragraph = _require_python_docx()
    _guard_archive_size(data, filename)

    try:
        document = docx.Document(io.BytesIO(data))
    except Exception as exc:
        # ZIP 가드는 통과했으나 python-docx가 문서를 해석하지 못한 경우다.
        # 사용자가 취할 조치는 ZIP 실패와 같으므로 같은 문구를 쓰고, 어느 단계에서
        # 실패했는지는 로그와 예외 체인으로 구분한다.
        logger.warning("DOCX 문서 파싱 실패 filename=%s", filename, exc_info=True)
        raise ConversionError(
            ErrorCode.CORRUPT_FILE,
            DOCX_OPEN_FAILED_MESSAGE,
            source=filename,
        ) from exc

    raw_blocks: list[dict] = []
    warnings: list[ConversionWarning] = []
    table_index = 0
    style_names = _style_name_map(document)

    for item in _iter_body(document, qn, Table, Paragraph):
        if isinstance(item, Table):
            table_blocks, flatten_warning, content_loss = _table_blocks(
                item, table_index, qn, Table, Paragraph
            )
            if table_blocks:
                raw_blocks.extend(table_blocks)
                warnings.append(flatten_warning)
            # 내용 유실 경고는 블록 유무와 무관하게 전달한다 — 깊이 상한 때문에
            # 표 전체가 비어 버린 경우가 정확히 "조용한 유실"이라 가장 알려야 한다.
            # (완전히 빈 표는 애초에 유실 경고를 만들지 않으므로 여기서 걸러지지 않는다)
            warnings.extend(content_loss)
            table_index += 1
            continue

        text = normalize_text(item.text).replace("\n", " ").strip()
        if not text or is_noise_line(text):
            continue
        kind, level = _paragraph_kind(item, style_names)
        raw_blocks.append({"kind": kind, "text": text, "level": level})

    # 이미지·차트·도형은 텍스트가 없어 그대로 유실된다. 조용히 버리지 않고 알린다.
    shape_count = _unsupported_element_count(document, qn)
    if shape_count:
        warnings.append(ConversionWarning(
            WarningCode.UNSUPPORTED_ELEMENT,
            f"이미지·도형·텍스트 상자 {shape_count}개는 텍스트로 변환되지 않았습니다.",
        ))

    document_out = assemble(filename, "docx", raw_blocks, warnings)
    document_out.blocks, dedupe_warnings = drop_duplicate_blocks(document_out.blocks)
    document_out.warnings.extend(dedupe_warnings)
    return document_out
